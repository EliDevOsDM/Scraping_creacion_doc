import flet as ft
import pandas as pd
import os
import asyncio
import numpy as np
import openpyxl
from pathlib import Path
from openpyxl.utils import get_column_letter, column_index_from_string
from openpyxl.styles import Border, Side, Font
from datetime import datetime
import warnings
from playwright.async_api import async_playwright

from app.config import WorkerConfig
from app.automation_service import AutomationService
from app.playwright_facade import PlaywrightFacade


def find_system_browser():
    """Encuentra Chrome o Edge instalado en el sistema"""
    browser_paths = [
        'C:\\Program Files\\Google\\Chrome\\Application\\chrome.exe',
        'C:\\Program Files (x86)\\Google\\Chrome\\Application\\chrome.exe',
        Path.home() / 'AppData\\Local\\Google\\Chrome\\Application\\chrome.exe',
        'C:\\Program Files (x86)\\Microsoft\\Edge\\Application\\msedge.exe',
        'C:\\Program Files\\Microsoft\\Edge\\Application\\msedge.exe',
    ]
    for path in browser_paths:
        browser_path = Path(path)
        if browser_path.exists():
            print(f'\u2713 Navegador encontrado: {browser_path}')
            return str(browser_path)
    return None


warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')


async def run_sgdea_automation(config_path: str, generados_dir: str, df_correos: pd.DataFrame,
                               page: ft.Page, status_text: ft.Text,
                               progress_bar: ft.ProgressBar) -> tuple[bool, str]:
    try:
        df_cred = pd.read_excel(config_path, sheet_name='Credenciales')
        user = str(df_cred.loc[0, 'Usuario'])
        password = str(df_cred.loc[0, 'Contraseña'])
    except Exception as e:
        return (False, f'Error obteniendo credenciales: {e}')

    base_dir = os.getcwd()
    task_screenshot_path = WorkerConfig.SCREENSHOT_PATH.value
    os.makedirs(WorkerConfig.DOWNLOAD_PATH.value, exist_ok=True)
    os.makedirs(task_screenshot_path, exist_ok=True)
    date_str = datetime.now().strftime('%d-%m-%Y')
    now_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    status_text.value = 'Iniciando sesión en SGDEA...'
    page.update()

    resultados = []
    try:
        async with async_playwright() as p:
            browser_config = WorkerConfig.random_browser_context()
            bot = await PlaywrightFacade.initialize(
                p,
                screenshot_path=task_screenshot_path,
                download_path=WorkerConfig.DOWNLOAD_PATH.value,
                headless=False,
                user_agent=WorkerConfig.random_user_agent(),
                viewport=browser_config.get('viewport'),
                locale=browser_config.get('locale'),
                timezone_id=browser_config.get('timezone_id'),
                screenshot_on_action=False,
            )
            service = AutomationService(bot)
            await service.login(user, password)
            status_text.value = 'Conectado a SGDEA. Procesando entidades...'
            page.update()
            total = len(df_correos)
            for idx, row in df_correos.iterrows():
                try:
                    data = row.to_dict()
                    entidad = data.get('Nombre entidad', '')
                    if pd.isna(entidad) or str(entidad).strip() == '':
                        continue
                    entidad = str(entidad).strip()
                    safe_entidad = ''.join(c for c in entidad if c.isalnum() or c in ' ._-')
                    excel_path = os.path.join(generados_dir, f'{safe_entidad} {date_str}.xlsx')
                    word_path = os.path.join(generados_dir, f'{safe_entidad} {date_str}.docx')
                    if not (os.path.exists(excel_path) and os.path.exists(word_path)):
                        print(f'Archivos no encontrados para {entidad}, saltando...')
                        continue
                    status_text.value = f'Subiendo y procesando {entidad} ({idx + 1}/{total})...'
                    page.update()
                    solicitud = await service.create_communication_correos(data, excel_path, word_path, base_dir)
                    resultados.append((excel_path, solicitud, now_str))
                except Exception as e:
                    print(f'Error automatizando {entidad}: {e}')
                    resultados.append((excel_path if 'excel_path' in locals() else 'N/A', f'Error: {e}', now_str))
                    continue
            await bot.close()
        if resultados:
            wb = openpyxl.load_workbook(config_path)
            if 'Historico' not in wb.sheetnames:
                if 'Historial' in wb.sheetnames:
                    ws = wb['Historial']
                else:
                    ws = wb.create_sheet('Historico')
                    ws.append(['Archivo', 'Solicitud', 'Fecha'])
            else:
                ws = wb['Historico']
            for res in resultados:
                ws.append(res)
            wb.save(config_path)
        return (True, 'Proceso automatizado en SGDEA completado.')
    except Exception as e:
        return (False, f'Error en automatización SGDEA: {e}')


async def main_app(page: ft.Page):
    page.title = 'Formato Decreto'
    page.vertical_alignment = ft.MainAxisAlignment.CENTER
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER
    page.window.width = 600
    page.window.height = 600
    page.window.resizable = False
    page.theme_mode = ft.ThemeMode.LIGHT
    await page.window.center()

    status_text = ft.Text("Procese archivos de entidades desde 'pendientes'", size=14,
                          text_align=ft.TextAlign.CENTER)
    progress_bar = ft.ProgressBar(width=400, visible=False)
    info_text = ft.Text('Carpeta origen: files/pendiente', size=12, color=ft.Colors.GREY_700)

    def process_entities_logic() -> tuple[bool, str]:
        try:
            import shutil
            from docx import Document
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt

            base_path = os.getcwd()
            config_path = os.path.join(base_path, 'files', 'config.xlsx')
            df_correos = pd.read_excel(config_path, sheet_name='Correos')
            df_zonas = pd.read_excel(config_path, sheet_name='Zonas')

            pendientes_dir = os.path.join(base_path, 'files', 'pendientes')
            if not os.path.exists(pendientes_dir):
                pendientes_dir = os.path.join(base_path, 'files', 'pendiente')

            plantilla_path = os.path.join(base_path, 'files', 'plantillas', 'Plantilla excel.xlsx')
            if not os.path.exists(plantilla_path):
                plantilla_path = os.path.join(base_path, 'files', 'plantillas', 'plantilla_excel.xlsx')

            generados_dir = os.path.join(base_path, 'files', 'generados')
            os.makedirs(generados_dir, exist_ok=True)

            if not os.path.exists(pendientes_dir):
                return (False, f'No existe la carpeta {pendientes_dir}')

            files_list = [f for f in os.listdir(pendientes_dir)
                          if f.endswith(('.xlsx', '.xls')) and not f.startswith('~')]
            if not files_list:
                return (False, 'No hay archivos en la carpeta de pendientes.')

            first_pendiente = os.path.join(pendientes_dir, files_list[0])
            try:
                df_headers = pd.read_excel(first_pendiente, sheet_name='CONSOLIDADO', nrows=0)
                col_names = df_headers.columns.tolist()
                col_types = {}
                if len(col_names) > 8:
                    col_types[col_names[8]] = str
                if len(col_names) > 9:
                    col_types[col_names[9]] = str
                df_consolidado = pd.read_excel(first_pendiente, sheet_name='CONSOLIDADO', dtype=col_types)
                df_consolidado.columns = df_consolidado.columns.astype(str).str.strip()
            except Exception as e:
                return (False, f"Error leyendo pestaña 'CONSOLIDADO': {e}")

            if 'DEPARTAMENTO' not in df_consolidado.columns:
                return (False, "La columna 'DEPARTAMENTO' no está en CONSOLIDADO")

            processed_entities = set()
            for idx, row in df_correos.iterrows():
                entidad = row.get('Nombre entidad')
                if pd.isna(entidad) or str(entidad).strip() == '':
                    continue
                entidad = str(entidad).strip()
                if entidad in processed_entities:
                    continue
                processed_entities.add(entidad)

                deptos_zonas = df_zonas[df_zonas['Nombre entidad'].astype(str).str.strip() == entidad]['Departamento'].dropna()
                deptos_list = [str(d).strip().upper() for d in deptos_zonas]
                if not deptos_list:
                    continue

                mask = df_consolidado['DEPARTAMENTO'].astype(str).str.strip().str.upper().isin(deptos_list)
                df_filtrado = df_consolidado[mask]
                if df_filtrado.empty:
                    continue

                try:
                    wb = openpyxl.load_workbook(plantilla_path)
                    ws = wb.active
                    from openpyxl.utils.dataframe import dataframe_to_rows
                    date_str = datetime.now().strftime('%d-%m-%Y')
                    for r_idx, r_data in enumerate(dataframe_to_rows(df_filtrado, index=False, header=False), 2):
                        for c_idx, value in enumerate(r_data, 1):
                            cell = ws.cell(row=r_idx, column=c_idx, value=value)
                            if c_idx in (9, 10):
                                cell.number_format = '@'
                        ws.cell(row=r_idx, column=21, value=str(entidad).strip())
                        ws.cell(row=r_idx, column=22, value='')
                        ws.cell(row=r_idx, column=35, value=date_str)
                    safe_entidad = ''.join(c for c in entidad if c.isalnum() or c in ' ._-')
                    out_filename = f'{safe_entidad} {date_str}.xlsx'
                    out_path = os.path.join(generados_dir, out_filename)
                    wb.save(out_path)
                except Exception as e:
                    print(f'Error procesando excel {entidad}: {e}')
                    continue

                try:
                    plantilla_word = os.path.join(base_path, 'files', 'plantillas', 'Plantilla word.docx')
                    if not os.path.exists(plantilla_word):
                        plantilla_word = os.path.join(base_path, 'files', 'plantillas', 'plantilla_word.docx')
                    word_filename = f'{safe_entidad} {date_str}.docx'
                    word_out_path = os.path.join(generados_dir, word_filename)
                    doc = Document(plantilla_word)

                    saludo_val = row.get('Saludos')
                    if pd.isna(saludo_val):
                        saludo_val = row.get('Saludo', '')
                    if str(saludo_val).strip() != '':
                        replace_str = str(saludo_val).strip()
                        for paragraph in doc.paragraphs:
                            if '[saludo]' in paragraph.text:
                                changed_runs = False
                                for run in paragraph.runs:
                                    if '[saludo]' in run.text:
                                        run.text = run.text.replace('[saludo]', replace_str)
                                        changed_runs = True
                                if not changed_runs:
                                    if '[saludo]' in paragraph.text:
                                        paragraph.text = paragraph.text.replace('[saludo]', replace_str)

                    target_paragraph = None
                    for paragraph in doc.paragraphs:
                        if 'casos:' in paragraph.text.lower():
                            target_paragraph = paragraph
                    if target_paragraph is None:
                        target_paragraph = doc.paragraphs[-1] if len(doc.paragraphs) > 0 else doc.add_paragraph()

                    p = target_paragraph._p
                    new_p1 = OxmlElement('w:p')
                    p.addnext(new_p1)
                    new_p2 = OxmlElement('w:p')
                    new_p1.addnext(new_p2)

                    num_rows = len(df_filtrado) + 1
                    table = doc.add_table(rows=num_rows, cols=5)
                    try:
                        table.style = 'Table Grid'
                    except KeyError:
                        try:
                            table.style = 'TableGrid'
                        except KeyError:
                            try:
                                table.style = 'Normal Table'
                            except KeyError:
                                pass
                    new_p2.addnext(table._tbl)

                    tbl = table._tbl
                    tblPr = tbl.tblPr
                    tblBorders = OxmlElement('w:tblBorders')
                    for b_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                        border = OxmlElement(f'w:{b_name}')
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '4')
                        border.set(qn('w:space'), '0')
                        border.set(qn('w:color'), 'auto')
                        tblBorders.append(border)
                    tblPr.append(tblBorders)

                    table.columns[0].width = Cm(1.5)
                    table.columns[4].width = Cm(4)

                    headers = ['No.', 'No. DE RADICADO MEN', 'DEMANDADO', 'NUMERO DE PROCESO. RAMA', 'ZONA']
                    for c_idx, header in enumerate(headers):
                        cell = table.cell(0, c_idx)
                        cell.text = header
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    for r_idx, (_, row_data) in enumerate(df_filtrado.iterrows(), start=1):
                        table.cell(r_idx, 0).text = str(r_idx)

                        col_rad = next((c for c in df_filtrado.columns if 'RADICAC' in c.upper()), None)
                        val_rad = row_data[col_rad] if col_rad in row_data.index else ''
                        table.cell(r_idx, 1).text = str(val_rad) if pd.notna(val_rad) and val_rad != '' else ''

                        col_dem = next((c for c in df_filtrado.columns if 'DEMANDADO' in c.upper()), None)
                        val_dem = row_data[col_dem] if col_dem in row_data.index else ''
                        table.cell(r_idx, 2).text = str(val_dem) if pd.notna(val_dem) and val_dem != '' else ''

                        col_proc = next((c for c in df_filtrado.columns if '23_DIGITOS' in c.upper()), None)
                        val_proc = row_data[col_proc] if col_proc in row_data.index else ''
                        table.cell(r_idx, 3).text = str(val_proc) if pd.notna(val_proc) and val_proc != '' else ''

                        col_zona = next((c for c in df_filtrado.columns if 'DEPARTAMENTO' in c.upper()), None)
                        val_zona = row_data[col_zona] if col_zona in row_data.index else ''
                        table.cell(r_idx, 4).text = str(val_zona) if pd.notna(val_zona) and val_zona != '' else ''

                    for row_tbl in table.rows:
                        for cell_tbl in row_tbl.cells:
                            for paragraph_tbl in cell_tbl.paragraphs:
                                for run_tbl in paragraph_tbl.runs:
                                    run_tbl.font.name = 'Verdana'
                                    run_tbl.font.size = Pt(11)

                    doc.save(word_out_path)
                except Exception as e:
                    print(f'Error procesando word {entidad}: {e}')
                    continue

            return (True, 'Los documentos de entidades se generaron correctamente.')
        except Exception as e:
            return (False, f'Error general: {str(e)}')

    async def on_click_entidades(e):
        btn_entidades.disabled = True
        progress_bar.visible = True
        status_text.value = 'Verificando navegador...'
        status_text.color = ft.Colors.BLUE
        page.update()

        browser_path = find_system_browser()
        if not browser_path:
            status_text.value = '\u274c Error: No se encontró Google Chrome ni Microsoft Edge.'
            status_text.color = ft.Colors.RED
            progress_bar.visible = False
            btn_entidades.disabled = False
            page.update()
            return

        os.environ['BROWSER_EXECUTABLE'] = browser_path
        status_text.value = 'Generando archivos por entidad...'
        page.update()

        success, message = await asyncio.to_thread(process_entities_logic)
        if success:
            status_text.value = f'\u2705 {message}. Iniciando automatización SGDEA...'
            page.update()
            base_path = os.getcwd()
            config_path = os.path.join(base_path, 'files', 'config.xlsx')
            generados_dir = os.path.join(base_path, 'files', 'generados')
            try:
                df_correos = pd.read_excel(config_path, sheet_name='Correos')
                sgdea_success, sgdea_message = await run_sgdea_automation(
                    config_path, generados_dir, df_correos, page, status_text, progress_bar)
                if sgdea_success:
                    status_text.value = f'\u2705 {sgdea_message}'
                    status_text.color = ft.Colors.GREEN
                else:
                    status_text.value = f'\u274c {message} pero falló SGDEA: {sgdea_message}'
                    status_text.color = ft.Colors.RED
            except Exception as ev:
                status_text.value = f'\u274c {message} pero falló lanzar SGDEA: {ev}'
                status_text.color = ft.Colors.RED
        else:
            status_text.value = f'\u274c Error: {message}'
            status_text.color = ft.Colors.RED

        progress_bar.visible = False
        btn_entidades.disabled = False
        page.update()

    btn_entidades = ft.Button(
        'Generar por Entidad',
        icon=ft.Icons.FILE_COPY,
        on_click=on_click_entidades,
        height=50,
        style=ft.ButtonStyle(color=ft.Colors.WHITE, bgcolor=ft.Colors.GREEN_700),
    )

    page.add(
        ft.Column(
            [
                ft.Icon(ft.Icons.TABLE_VIEW, size=64, color=ft.Colors.BLUE),
                ft.Text('Formato Decreto', size=24, weight=ft.FontWeight.BOLD),
                ft.Container(height=20),
                info_text,
                ft.Container(height=10),
                ft.Row([btn_entidades], alignment=ft.MainAxisAlignment.CENTER),
                ft.Container(height=20),
                progress_bar,
                ft.Container(height=10),
                status_text,
            ],
            alignment=ft.MainAxisAlignment.CENTER,
            horizontal_alignment=ft.CrossAxisAlignment.CENTER,
        )
    )


if __name__ == '__main__':
    ft.run(main_app)
